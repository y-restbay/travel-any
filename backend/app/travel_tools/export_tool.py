"""行程导出工具:把已生成的 itinerary 渲染成 Word / PDF。

设计要点:
- 磁盘文件名固定为 ``{itinerary_id}.{format}``,避免中文 / 路径穿越问题
- 展示用 ``display_filename`` 在 SSE 里推送,前端做 a[download] 属性
- 文件已存在则直接复用(LLM 反复说"导出 PDF"时不重新生成)
- 中文字体:
  * PDF 用 reportlab 注册思源黑体;字体文件缺失时 fallback Helvetica 并打日志
  * Word 通过 docx oxml 设置 eastAsia 字体为宋体(系统都自带)
- 静态地图截图:预留 hook 但默认关闭(避免在没有 AMAP_KEY 时阻塞主路径)
"""
from __future__ import annotations

import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from app.travel_tools.itinerary_store import get_itinerary

logger = logging.getLogger(__name__)


# --------------------------------------------------------------------- 路径与字体
_BACKEND_DIR = Path(__file__).resolve().parents[2]  # .../backend
EXPORT_DIR = _BACKEND_DIR / "storage" / "exports"
FONT_PATH = _BACKEND_DIR / "assets" / "fonts" / "SourceHanSansCN-Regular.otf"

_PDF_FONT_NAME = "SourceHan"
_PDF_FONT_REGISTERED: Optional[bool] = None  # None=未尝试, True=注册成功, False=失败已 fallback


def _ensure_pdf_font_registered() -> str:
    """惰性注册思源黑体。返回最终可用的字体名(Helvetica 是 fallback)。"""
    global _PDF_FONT_REGISTERED
    if _PDF_FONT_REGISTERED is True:
        return _PDF_FONT_NAME
    if _PDF_FONT_REGISTERED is False:
        return "Helvetica"

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    if FONT_PATH.exists():
        try:
            pdfmetrics.registerFont(TTFont(_PDF_FONT_NAME, str(FONT_PATH)))
            _PDF_FONT_REGISTERED = True
            return _PDF_FONT_NAME
        except Exception as exc:
            logger.warning("注册思源黑体失败,退化为 Helvetica:%s", exc)
            _PDF_FONT_REGISTERED = False
            return "Helvetica"
    logger.warning(
        "未找到中文字体 %s,PDF 中文将变方块。请把 SourceHanSansCN-Regular.otf 放到该路径。",
        FONT_PATH,
    )
    _PDF_FONT_REGISTERED = False
    return "Helvetica"


# --------------------------------------------------------------------- Tool Schema
EXPORT_TOOL_SCHEMA: Dict[str, Any] = {
    "type": "function",
    "function": {
        "name": "export_itinerary",
        "description": (
            "把已经通过 generate_itinerary_summary 生成的行程导出为可下载的 PDF 或 Word 文件。"
            "调用前提:用户已经看到行程卡片,并明确表达了导出意愿。"
            "导出完成后会通过 SSE 推送下载链接,你只需要简短确认即可。"
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "itinerary_id": {
                    "type": "string",
                    "description": "由 generate_itinerary_summary 返回的 itin_xxx,必填。",
                },
                "format": {
                    "type": "string",
                    "enum": ["pdf", "docx"],
                    "description": "导出格式,默认 pdf。用户说 'Word' / '文档' 则用 docx。",
                    "default": "pdf",
                },
                "include_map_snapshot": {
                    "type": "boolean",
                    "description": "是否在文档中嵌入高德静态地图截图。默认 true,但当前实现暂未启用,会被忽略。",
                    "default": True,
                },
            },
            "required": ["itinerary_id"],
        },
    },
}


# --------------------------------------------------------------------- 主入口
def handle_export_itinerary(
    args: Dict[str, Any],
    *,
    session_id: str = "",
) -> Tuple[Dict[str, Any], Optional[Dict[str, Any]]]:
    """导出行程。返回 ``(summary_for_llm, sse_payload_or_none)``。"""
    if not isinstance(args, dict):
        return {"error": "参数必须是对象。"}, None

    itinerary_id = (args.get("itinerary_id") or "").strip()
    fmt = (args.get("format") or "pdf").lower().strip()
    if fmt not in {"pdf", "docx"}:
        return {"error": f"format 必须是 pdf 或 docx,得到 '{fmt}'。"}, None
    if not itinerary_id:
        return {"error": "缺少 itinerary_id;请先调用 generate_itinerary_summary。"}, None

    itinerary = get_itinerary(itinerary_id)
    if itinerary is None:
        return (
            {
                "error": f"找不到 itinerary_id={itinerary_id};可能尚未生成行程,或服务已重启。",
                "提示": "请先调用 generate_itinerary_summary 生成行程,再导出。",
            },
            None,
        )

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    file_path = EXPORT_DIR / f"{itinerary_id}.{fmt}"

    try:
        if not file_path.exists():
            if fmt == "docx":
                _export_docx(itinerary, file_path)
            else:
                _export_pdf(itinerary, file_path)
    except Exception as exc:
        logger.exception("导出失败:%s", exc)
        return {"error": f"导出失败:{exc.__class__.__name__}: {exc}"}, None

    size_bytes = file_path.stat().st_size
    size_text = _format_size(size_bytes)
    display_filename = f"{_safe_filename(itinerary.get('trip_title') or 'itinerary')}_{itinerary_id}.{fmt}"
    download_url = f"/api/exports/{itinerary_id}.{fmt}"

    sse_payload: Dict[str, Any] = {
        "type": "export_ready",
        "itinerary_id": itinerary_id,
        "format": fmt,
        "filename": display_filename,
        "download_url": download_url,
        "size_bytes": size_bytes,
        "size_text": size_text,
        "trip_title": itinerary.get("trip_title", ""),
    }
    summary_for_llm = {
        "状态": "success",
        "文件格式": fmt.upper(),
        "文件名": display_filename,
        "下载链接": download_url,
        "文件大小": size_text,
        "提示": "导出完成,下载链接已发送给用户,请简短确认即可。",
    }
    return summary_for_llm, sse_payload


# --------------------------------------------------------------------- Word 渲染
def _export_docx(itinerary: Dict[str, Any], path: Path) -> None:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document()
    _set_doc_default_chinese_font(doc, "宋体")

    # ---- 封面 ----
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(itinerary.get("trip_title") or "我的行程")
    run.font.size = Pt(28)
    run.bold = True

    if itinerary.get("trip_dates"):
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run(itinerary["trip_dates"])
        sub_run.font.size = Pt(14)

    gen = doc.add_paragraph()
    gen.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    gen_run = gen.add_run(f"生成于 {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    gen_run.font.size = Pt(9)

    doc.add_page_break()

    # ---- 概览 ----
    _add_heading(doc, "行程概览", 1)
    if itinerary.get("summary"):
        doc.add_paragraph(itinerary["summary"])

    meta = itinerary.get("meta") or {}
    if meta:
        for label_zh, key in [
            ("目的地", "destination"),
            ("人数", "people"),
            ("预算", "budget"),
            ("住宿", "accommodation"),
            ("偏好", "preferences"),
            ("主要出行方式", "transport_mode"),
        ]:
            value = meta.get(key)
            if value:
                p = doc.add_paragraph()
                p.add_run(f"{label_zh}:").bold = True
                p.add_run(str(value))

    weather = itinerary.get("weather_summary") or []
    if weather:
        _add_heading(doc, "天气概览", 2)
        for item in weather:
            line = " · ".join(
                str(item.get(k))
                for k in ("date", "condition", "temp", "tip")
                if item.get(k)
            )
            if line:
                doc.add_paragraph(line, style="List Bullet")

    notes = itinerary.get("important_notes") or []
    if notes:
        _add_heading(doc, "出行须知", 2)
        for note in notes:
            doc.add_paragraph(str(note), style="List Number")

    # ---- 每日安排 ----
    for day in itinerary.get("days") or []:
        doc.add_page_break()
        day_num = day.get("day_number", "")
        day_title = day.get("title", "")
        _add_heading(doc, f"Day {day_num} · {day_title}", 1)
        if day.get("theme"):
            doc.add_paragraph(day["theme"])

        schedule = day.get("schedule") or []
        if schedule:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "时间"
            hdr[1].text = "类型"
            hdr[2].text = "地点 / 说明"
            hdr[3].text = "时长 / 花费"
            for item in schedule:
                row = table.add_row().cells
                row[0].text = str(item.get("time", ""))
                row[1].text = _schedule_type_label(item.get("type", ""))
                row[2].text = _format_schedule_detail(item)
                row[3].text = _format_schedule_cost(item)
                for cell in row:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        day_cost = day.get("day_cost") or {}
        if day_cost:
            p = doc.add_paragraph()
            p.add_run("当日花费:").bold = True
            p.add_run(_format_day_cost(day_cost))

    # ---- 总预算 ----
    total = itinerary.get("total_budget") or {}
    if total:
        doc.add_page_break()
        _add_heading(doc, "总预算汇总", 1)
        budget_table = doc.add_table(rows=1, cols=2)
        budget_table.style = "Light Grid Accent 1"
        budget_table.rows[0].cells[0].text = "类别"
        budget_table.rows[0].cells[1].text = "金额 (元)"
        for label_zh, key in [
            ("门票", "tickets"),
            ("餐饮", "meals"),
            ("交通", "transport"),
            ("住宿", "accommodation"),
            ("合计", "total"),
        ]:
            value = total.get(key)
            if value is None:
                continue
            row = budget_table.add_row().cells
            row[0].text = label_zh
            row[1].text = str(value)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_run = foot.add_run("由 WanderBot 漫游指南 生成")
    foot_run.font.size = Pt(9)

    doc.save(str(path))


def _set_doc_default_chinese_font(doc, font_name: str) -> None:
    """让 Word 文档的默认样式同时支持西文与东亚字体。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def _add_heading(doc, text: str, level: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "宋体"
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), "宋体")


# --------------------------------------------------------------------- PDF 渲染
def _export_pdf(itinerary: Dict[str, Any], path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    font_name = _ensure_pdf_font_registered()
    title_style = ParagraphStyle(
        "title", fontName=font_name, fontSize=26, leading=32, alignment=TA_CENTER, spaceAfter=18
    )
    subtitle_style = ParagraphStyle(
        "sub", fontName=font_name, fontSize=14, leading=20, alignment=TA_CENTER, spaceAfter=12
    )
    gen_style = ParagraphStyle(
        "gen", fontName=font_name, fontSize=9, leading=12, alignment=TA_RIGHT, textColor=colors.grey
    )
    h1 = ParagraphStyle("h1", fontName=font_name, fontSize=18, leading=24, spaceAfter=10)
    h2 = ParagraphStyle("h2", fontName=font_name, fontSize=14, leading=20, spaceAfter=8, textColor=colors.HexColor("#65735D"))
    body = ParagraphStyle("body", fontName=font_name, fontSize=11, leading=17, alignment=TA_LEFT, spaceAfter=4)

    story: List[Any] = []

    # 封面
    story.append(Spacer(1, 6 * cm))
    story.append(Paragraph(_xml_escape(itinerary.get("trip_title") or "我的行程"), title_style))
    if itinerary.get("trip_dates"):
        story.append(Paragraph(_xml_escape(itinerary["trip_dates"]), subtitle_style))
    story.append(Spacer(1, 4 * cm))
    story.append(Paragraph(f"生成于 {datetime.now().strftime('%Y-%m-%d %H:%M')}", gen_style))
    story.append(PageBreak())

    # 概览
    story.append(Paragraph("行程概览", h1))
    if itinerary.get("summary"):
        story.append(Paragraph(_xml_escape(itinerary["summary"]), body))

    meta = itinerary.get("meta") or {}
    if meta:
        for label_zh, key in [
            ("目的地", "destination"),
            ("人数", "people"),
            ("预算", "budget"),
            ("住宿", "accommodation"),
            ("偏好", "preferences"),
            ("主要出行方式", "transport_mode"),
        ]:
            value = meta.get(key)
            if value:
                story.append(Paragraph(f"<b>{label_zh}:</b> {_xml_escape(str(value))}", body))

    weather = itinerary.get("weather_summary") or []
    if weather:
        story.append(Spacer(1, 0.4 * cm))
        story.append(Paragraph("天气概览", h2))
        for item in weather:
            line = " · ".join(
                str(item.get(k))
                for k in ("date", "condition", "temp", "tip")
                if item.get(k)
            )
            if line:
                story.append(Paragraph("• " + _xml_escape(line), body))

    notes = itinerary.get("important_notes") or []
    if notes:
        story.append(Spacer(1, 0.4 * cm))
        story.append(Paragraph("出行须知", h2))
        for i, note in enumerate(notes, start=1):
            story.append(Paragraph(f"{i}. {_xml_escape(str(note))}", body))

    # 每日
    for day in itinerary.get("days") or []:
        story.append(PageBreak())
        title_line = f"Day {day.get('day_number', '')} · {_xml_escape(day.get('title') or '')}"
        story.append(Paragraph(title_line, h1))
        if day.get("theme"):
            story.append(Paragraph(_xml_escape(day["theme"]), body))

        schedule = day.get("schedule") or []
        if schedule:
            table_data: List[List[Any]] = [["时间", "类型", "地点 / 说明", "时长 / 花费"]]
            for item in schedule:
                table_data.append(
                    [
                        Paragraph(_xml_escape(str(item.get("time", ""))), body),
                        Paragraph(_xml_escape(_schedule_type_label(item.get("type", ""))), body),
                        Paragraph(_xml_escape(_format_schedule_detail(item)), body),
                        Paragraph(_xml_escape(_format_schedule_cost(item)), body),
                    ]
                )
            table = Table(table_data, colWidths=[2.5 * cm, 2 * cm, 8 * cm, 3.5 * cm], repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2EAE0")),
                        ("FONT", (0, 0), (-1, -1), font_name, 10),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D9D3CA")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )
            story.append(Spacer(1, 0.3 * cm))
            story.append(table)

        day_cost = day.get("day_cost") or {}
        if day_cost:
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph(f"<b>当日花费:</b> {_xml_escape(_format_day_cost(day_cost))}", body))

    # 总预算
    total = itinerary.get("total_budget") or {}
    if total:
        story.append(PageBreak())
        story.append(Paragraph("总预算汇总", h1))
        budget_data = [["类别", "金额 (元)"]]
        for label_zh, key in [
            ("门票", "tickets"),
            ("餐饮", "meals"),
            ("交通", "transport"),
            ("住宿", "accommodation"),
            ("合计", "total"),
        ]:
            value = total.get(key)
            if value is None:
                continue
            budget_data.append([label_zh, str(value)])
        if len(budget_data) > 1:
            budget_table = Table(budget_data, colWidths=[6 * cm, 6 * cm])
            budget_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2EAE0")),
                        ("FONT", (0, 0), (-1, -1), font_name, 11),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D9D3CA")),
                        ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                    ]
                )
            )
            story.append(budget_table)

    story.append(Spacer(1, 1 * cm))
    story.append(Paragraph("由 WanderBot 漫游指南 生成", gen_style))

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=itinerary.get("trip_title") or "WanderBot Itinerary",
    )
    doc.build(story)


# --------------------------------------------------------------------- 格式化辅助
def _schedule_type_label(t: str) -> str:
    return {
        "depart": "出发",
        "visit": "游览",
        "meal": "用餐",
        "transit": "中转",
        "return": "返程",
    }.get((t or "").lower(), t or "")


def _format_schedule_detail(item: Dict[str, Any]) -> str:
    parts: List[str] = []
    place = item.get("place")
    if place:
        parts.append(str(place))
    elif item.get("from") or item.get("to"):
        parts.append(f"{item.get('from', '')} → {item.get('to', '')}")
    if item.get("note"):
        parts.append(str(item["note"]))
    if item.get("highlights"):
        parts.append("亮点:" + _stringify_list(item["highlights"]))
    if item.get("must_try"):
        parts.append("推荐:" + _stringify_list(item["must_try"]))
    if item.get("cuisine"):
        parts.append("菜系:" + str(item["cuisine"]))
    if item.get("tips"):
        parts.append("贴士:" + str(item["tips"]))
    return "\n".join(parts) if parts else ""


def _format_schedule_cost(item: Dict[str, Any]) -> str:
    parts: List[str] = []
    if item.get("duration_min"):
        parts.append(f"{item['duration_min']} 分钟")
    if item.get("ticket"):
        parts.append(f"门票 ¥{item['ticket']}")
    if item.get("cost"):
        parts.append(f"花费 ¥{item['cost']}")
    return " / ".join(parts)


def _format_day_cost(cost: Dict[str, Any]) -> str:
    pieces: List[str] = []
    for label_zh, key in [
        ("门票", "tickets"),
        ("餐饮", "meals"),
        ("交通", "transport"),
        ("合计", "total"),
    ]:
        value = cost.get(key)
        if value is not None:
            pieces.append(f"{label_zh} ¥{value}")
    return " / ".join(pieces)


def _stringify_list(value: Any) -> str:
    if isinstance(value, list):
        return "、".join(str(v) for v in value)
    return str(value)


def _xml_escape(text: str) -> str:
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


_FILENAME_BAD_CHARS = re.compile(r"[^\w一-龥\-]+")


def _safe_filename(name: str) -> str:
    cleaned = _FILENAME_BAD_CHARS.sub("_", name).strip("_")
    return cleaned[:40] or "itinerary"


def _format_size(num_bytes: int) -> str:
    if num_bytes >= 1024 * 1024:
        return f"{num_bytes / 1024 / 1024:.2f} MB"
    if num_bytes >= 1024:
        return f"{num_bytes / 1024:.1f} KB"
    return f"{num_bytes} B"
